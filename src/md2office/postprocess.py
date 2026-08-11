"""Post-process a pandoc-generated DOCX to match SSC book output.

Applied after pandoc emits the DOCX (with the SSC reference template):

1. Header injection — title + bilingual classification in every page header
   (and in the docProps/core.xml title).  Replaces the template's
   "[Enter Document Title]" placeholder and any stale metadata title.
2. Table width/style fixes — pandoc tables are fixed-width (~3.5 in) which
   LibreOffice crushes; rewrite to 100% pct width, autofit layout, explicit
   borders, and remap/inject the missing Compact/Table/Title/Subtitle styles.
3. Heading2 fix — drop keepLines + outlineLvl (a LibreOffice wrap bug).
4. Classification text box — widen the header sensitivity label box and set an
   explicit font size so it fits on one line.
5. Title-page heading indent — keep Heading1 on the cover away from the
   decorative leaf graphic.
6. Code block style — SourceCode paragraphs + VerbatimChar character style get
   a monospace font, 9pt size, and light-grey shading.
7. Section page breaks — optional; each Heading1 starts a new page.
8. Font switch — optional; switch document fonts to Arial.
9. Table font compaction — table cell runs are set to 9pt (python-docx pass).
"""

from __future__ import annotations

import os
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET
from xml.sax.saxutils import escape as xml_escape

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A14_NS = "http://schemas.microsoft.com/office/drawing/2010/main"

W = "{%s}" % W_NS
A = "{%s}" % A_NS
WP = "{%s}" % WP_NS

_NS_MAP = {
    "wpc": "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
    "cx": "http://schemas.microsoft.com/office/drawing/2014/chartex",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "aink": "http://schemas.microsoft.com/office/drawing/2016/ink",
    "am3d": "http://schemas.microsoft.com/office/drawing/2017/model3d",
    "o": "urn:schemas-microsoft-com:office:office",
    "oel": "http://schemas.microsoft.com/office/2019/extlst",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "v": "urn:schemas-microsoft-com:vml",
    "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
    "wp": WP_NS,
    "w10": "urn:schemas-microsoft-com:office:word",
    "w": W_NS,
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
    "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
    "w16cid": "http://schemas.microsoft.com/office/word/2016/wordml/cid",
    "w16": "http://schemas.microsoft.com/office/word/2018/wordml",
    "w16se": "http://schemas.microsoft.com/office/word/2015/wordml/symex",
    "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    "wpi": "http://schemas.microsoft.com/office/word/2010/wordprocessingInk",
    "wne": "http://schemas.microsoft.com/office/word/2006/wordml",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "a": A_NS,
    "a14": A14_NS,
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "adec": "http://schemas.microsoft.com/office/drawing/2017/decorative",
    "asvg": "http://schemas.microsoft.com/office/drawing/2016/SVG/main",
    "aclsh": "http://schemas.microsoft.com/office/drawing/2020/classificationShape",
}
for _prefix, _uri in _NS_MAP.items():
    ET.register_namespace(_prefix, _uri)

PAGE_PCT = "5000"
GRID_TWIPS = 9000
CLASSIFICATION_BOX_EMU = "1800000"
CLASSIFICATION_FONT_SZ = "16"
TITLEPAGE_IND_RIGHT = "3600"
TBLSTYLE_REMAP = {"Table": "TableNormal"}
CODE_BG_FILL = "F5F5F5"
CODE_FONT_SZ = "18"
CODE_FONT = "Consolas"

_DEFAULT_CLASSIFICATION = "Unclassified | Non classifié"


def _make_border(tag: str, val: str = "single", sz: str = "4",
                 space: str = "0", color: str = "auto") -> ET.Element:
    el = ET.Element(W + tag)
    el.set(W + "val", val)
    el.set(W + "sz", sz)
    el.set(W + "space", space)
    el.set(W + "color", color)
    return el


def _fix_table(tbl: ET.Element) -> None:
    tblPr = tbl.find(W + "tblPr")
    if tblPr is None:
        tblPr = ET.SubElement(tbl, W + "tblPr")

    tblStyle = tblPr.find(W + "tblStyle")
    if tblStyle is not None:
        current = tblStyle.get(W + "val", "")
        if current in TBLSTYLE_REMAP:
            tblStyle.set(W + "val", TBLSTYLE_REMAP[current])

    tblW = tblPr.find(W + "tblW")
    if tblW is None:
        tblW = ET.SubElement(tblPr, W + "tblW")
    tblW.set(W + "type", "pct")
    tblW.set(W + "w", PAGE_PCT)

    tblLayout = tblPr.find(W + "tblLayout")
    if tblLayout is None:
        tblLayout = ET.SubElement(tblPr, W + "tblLayout")
    tblLayout.set(W + "type", "autofit")

    tblBorders = tblPr.find(W + "tblBorders")
    if tblBorders is None:
        tblBorders = ET.SubElement(tblPr, W + "tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if tblBorders.find(W + side) is None:
            tblBorders.append(_make_border(side))

    tblGrid = tbl.find(W + "tblGrid")
    if tblGrid is not None:
        cols = tblGrid.findall(W + "gridCol")
        if cols:
            per = GRID_TWIPS // len(cols)
            for c in cols:
                c.set(W + "w", str(per))


def _fix_heading2_styles(styles_root: ET.Element) -> None:
    for st in styles_root.iter(W + "style"):
        if st.get(W + "styleId") == "Heading2":
            pPr = st.find(W + "pPr")
            if pPr is not None:
                for child in list(pPr):
                    if child.tag in (W + "keepLines", W + "outlineLvl"):
                        pPr.remove(child)


def _fix_missing_styles(styles_root: ET.Element) -> list:
    existing = {st.get(W + "styleId") for st in styles_root.iter(W + "style")}
    injected = []

    if "Compact" not in existing:
        st = ET.SubElement(styles_root, W + "style")
        st.set(W + "type", "paragraph")
        st.set(W + "styleId", "Compact")
        ET.SubElement(st, W + "name").set(W + "val", "Compact")
        based = ET.SubElement(st, W + "basedOn")
        based.set(W + "val", "Normal")
        pPr = ET.SubElement(st, W + "pPr")
        spacing = ET.SubElement(pPr, W + "spacing")
        spacing.set(W + "after", "0")
        spacing.set(W + "line", "240")
        spacing.set(W + "lineRule", "auto")
        injected.append("Compact")

    if "Table" not in existing and "TableNormal" in existing:
        st = ET.SubElement(styles_root, W + "style")
        st.set(W + "type", "table")
        st.set(W + "styleId", "Table")
        ET.SubElement(st, W + "name").set(W + "val", "Table")
        based = ET.SubElement(st, W + "basedOn")
        based.set(W + "val", "TableNormal")
        injected.append("Table")

    for st in styles_root.iter(W + "style"):
        if st.get(W + "styleId") in ("Title", "Subtitle"):
            pPr = st.find(W + "pPr")
            if pPr is None:
                pPr = ET.SubElement(st, W + "pPr")
            ind = pPr.find(W + "ind")
            if ind is None:
                ind = ET.SubElement(pPr, W + "ind")
            ind.set(W + "right", "3600")

    return injected


def _fix_classification_textboxes(root: ET.Element) -> int:
    fixed = 0
    for anchor in root.iter(WP + "anchor"):
        docPr = anchor.find(WP + "docPr")
        if docPr is None:
            continue
        descr = (docPr.get("descr") or "").lower()
        name = (docPr.get("name") or "").lower()
        if "classif" not in descr and "classif" not in name and "sensitivity" not in name:
            continue
        extent = anchor.find(WP + "extent")
        if extent is not None:
            extent.set("cx", CLASSIFICATION_BOX_EMU)
            fixed += 1
        for xfrm in anchor.iter(A + "xfrm"):
            ext = xfrm.find(A + "ext")
            if ext is not None:
                ext.set("cx", CLASSIFICATION_BOX_EMU)
        for rPr in anchor.iter(W + "rPr"):
            for tag in (W + "sz", W + "szCs"):
                el = rPr.find(tag)
                if el is None:
                    el = ET.SubElement(rPr, tag)
                el.set(W + "val", CLASSIFICATION_FONT_SZ)
    return fixed


def _fix_titlepage_headings(doc_root: ET.Element) -> int:
    fixed = 0
    for p in doc_root.iter(W + "p"):
        for br in p.iter(W + "br"):
            if br.get(W + "type") in ("page", "column"):
                return fixed
        pPr = p.find(W + "pPr")
        if pPr is None:
            continue
        pStyle = pPr.find(W + "pStyle")
        if pStyle is None or pStyle.get(W + "val") != "Heading1":
            continue
        ind = pPr.find(W + "ind")
        if ind is None:
            ind = ET.SubElement(pPr, W + "ind")
        current = int(ind.get(W + "right", "0") or "0")
        if current < int(TITLEPAGE_IND_RIGHT):
            ind.set(W + "right", TITLEPAGE_IND_RIGHT)
            fixed += 1
    return fixed


def _fix_code_block_style(styles_root: ET.Element) -> bool:
    fixed = False
    for st in styles_root.iter(W + "style"):
        if st.get(W + "styleId") == "SourceCode":
            pPr = st.find(W + "pPr")
            if pPr is None:
                pPr = ET.SubElement(st, W + "pPr")
            shd = pPr.find(W + "shd")
            if shd is None:
                shd = ET.SubElement(pPr, W + "shd")
            shd.set(W + "val", "clear")
            shd.set(W + "color", "auto")
            shd.set(W + "fill", CODE_BG_FILL)
            rPr = st.find(W + "rPr")
            if rPr is None:
                rPr = ET.SubElement(st, W + "rPr")
            rFonts = rPr.find(W + "rFonts")
            if rFonts is None:
                rFonts = ET.SubElement(rPr, W + "rFonts")
            for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                rFonts.set(W + attr, CODE_FONT)
            for tag in ("sz", "szCs"):
                el = rPr.find(W + tag)
                if el is None:
                    el = ET.SubElement(rPr, W + tag)
                el.set(W + "val", CODE_FONT_SZ)
            fixed = True
        if st.get(W + "styleId") == "VerbatimChar":
            rPr = st.find(W + "rPr")
            if rPr is None:
                rPr = ET.SubElement(st, W + "rPr")
            rFonts = rPr.find(W + "rFonts")
            if rFonts is None:
                rFonts = ET.SubElement(rPr, W + "rFonts")
            for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                rFonts.set(W + attr, CODE_FONT)
            for tag in ("sz", "szCs"):
                el = rPr.find(W + tag)
                if el is None:
                    el = ET.SubElement(rPr, W + tag)
                el.set(W + "val", CODE_FONT_SZ)
            fixed = True
    if not any(st.get(W + "styleId") == "VerbatimChar"
               for st in styles_root.iter(W + "style")):
        st = ET.SubElement(styles_root, W + "style")
        st.set(W + "type", "character")
        st.set(W + "styleId", "VerbatimChar")
        ET.SubElement(st, W + "name").set(W + "val", "VerbatimChar")
        based = ET.SubElement(st, W + "basedOn")
        based.set(W + "val", "DefaultParagraphFont")
        rPr = ET.SubElement(st, W + "rPr")
        rFonts = ET.SubElement(rPr, W + "rFonts")
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            rFonts.set(W + attr, CODE_FONT)
        for tag in ("sz", "szCs"):
            el = ET.SubElement(rPr, W + tag)
            el.set(W + "val", CODE_FONT_SZ)
        fixed = True
    return fixed


def _fix_font_arial(styles_root: ET.Element) -> int:
    fixed = 0
    docDefaults = styles_root.find(W + "docDefaults")
    if docDefaults is None:
        docDefaults = ET.SubElement(styles_root, W + "docDefaults")
    rPrDefault = docDefaults.find(W + "rPrDefault")
    if rPrDefault is None:
        rPrDefault = ET.SubElement(docDefaults, W + "rPrDefault")
    rPr = rPrDefault.find(W + "rPr")
    if rPr is None:
        rPr = ET.SubElement(rPrDefault, W + "rPr")
    rFonts = rPr.find(W + "rFonts")
    if rFonts is None:
        rFonts = ET.SubElement(rPr, W + "rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(W + attr, "Arial")
    fixed += 1

    MONO = {"SourceCode", "VerbatimChar"}
    for st in styles_root.iter(W + "style"):
        style_id = st.get(W + "styleId", "")
        rPr = st.find(W + "rPr")
        if rPr is None:
            rPr = ET.SubElement(st, W + "rPr")
        rFonts = rPr.find(W + "rFonts")
        if rFonts is None:
            rFonts = ET.SubElement(rPr, W + "rFonts")
        if style_id in MONO or style_id.endswith("Tok"):
            rFonts.set(W + "ascii", "Inconsolata")
            rFonts.set(W + "hAnsi", "Inconsolata")
            rFonts.set(W + "cs", "Inconsolata")
            continue
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            rFonts.set(W + attr, "Arial")
        fixed += 1
    return fixed


def _fix_section_page_breaks(doc_root: ET.Element) -> int:
    fixed = 0
    for p in doc_root.iter(W + "p"):
        pPr = p.find(W + "pPr")
        if pPr is None:
            continue
        pStyle = pPr.find(W + "pStyle")
        if pStyle is None or pStyle.get(W + "val") != "Heading1":
            continue
        if pPr.find(W + "pageBreakBefore") is not None:
            continue
        ET.SubElement(pPr, W + "pageBreakBefore")
        fixed += 1
    return fixed


def _rewrite_zip(path: Path, transform) -> dict:
    tmp = str(path) + ".tmp"
    counts = {"tables": 0, "classif_boxes": 0, "styles": False,
              "injected": [], "code_style": False, "page_breaks": 0,
              "font": 0}
    with zipfile.ZipFile(path) as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for name in zin.namelist():
                data = zin.read(name)
                if name == "word/styles.xml":
                    root = ET.fromstring(data)
                    _fix_heading2_styles(root)
                    counts["injected"] = _fix_missing_styles(root)
                    if transform["code_style"]:
                        counts["code_style"] = _fix_code_block_style(root)
                    if transform["font"] == "arial":
                        counts["font"] = _fix_font_arial(root)
                    counts["styles"] = True
                    data = ET.tostring(root, xml_declaration=True, encoding="UTF-8")
                elif name.startswith("word/header") and name.endswith(".xml"):
                    root = ET.fromstring(data)
                    counts["classif_boxes"] += _fix_classification_textboxes(root)
                    data = ET.tostring(root, xml_declaration=True, encoding="UTF-8")
                elif name.startswith("word/") and name.endswith(".xml") and b"<w:tbl>" in data:
                    root = ET.fromstring(data)
                    for tbl in root.iter(W + "tbl"):
                        _fix_table(tbl)
                        counts["tables"] += 1
                    if transform["page_breaks"] == "sections":
                        counts["page_breaks"] = _fix_section_page_breaks(root)
                    data = ET.tostring(root, xml_declaration=True, encoding="UTF-8")
                zout.writestr(name, data)
    os.replace(tmp, path)
    return counts


def _update_header(docx_path: Path, title: str, classification: str) -> None:
    """Replace the template header placeholder and core.xml title."""
    if not docx_path.exists():
        raise FileNotFoundError(f"Cannot find DOCX file at '{docx_path}'")

    title_xml = xml_escape(title)
    classification_xml = xml_escape(classification)

    temp_docx = str(docx_path) + ".tmp"
    shutil.copy2(docx_path, temp_docx)
    try:
        with zipfile.ZipFile(temp_docx, "r") as z:
            with tempfile.TemporaryDirectory() as tmpdir:
                z.extractall(tmpdir)

                core_xml_path = os.path.join(tmpdir, "docProps/core.xml")
                metadata_title = None
                if os.path.exists(core_xml_path):
                    with open(core_xml_path, "r", encoding="utf-8", errors="replace") as f:
                        content = f.read()
                    match = re.search(r"<dc:title>([^<]*)</dc:title>", content)
                    if match:
                        metadata_title = match.group(1)

                placeholders = ["[Enter Document Title]", "[Enter ",
                                "Document Title", "]"]
                if metadata_title and metadata_title.strip():
                    placeholders.append(metadata_title)
                seen = set()
                placeholders = [p for p in placeholders if p and p.strip()
                                and not (p in seen or seen.add(p))]

                if os.path.exists(core_xml_path):
                    with open(core_xml_path, "r", encoding="utf-8", errors="replace") as f:
                        content = f.read()
                    content = re.sub(
                        r"<dc:title>([^<]*)</dc:title>",
                        f"<dc:title>{title_xml}</dc:title>",
                        content,
                    )
                    with open(core_xml_path, "w", encoding="utf-8") as f:
                        f.write(content)

                for name in z.namelist():
                    if "header" in name.lower() and name.endswith(".xml"):
                        filepath = os.path.join(tmpdir, name)
                        if not os.path.exists(filepath):
                            continue
                        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                            content = f.read()
                        split_pattern = r"\[Enter [\s\S]*?Document Title[\s\S]*?\]"
                        split_replacement = (
                            f"<w:r><w:t>{title_xml}</w:t></w:r>"
                        )
                        content = re.sub(split_pattern, split_replacement,
                                         content, count=1)
                        for placeholder in placeholders:
                            if placeholder in ("[Enter ", "Document Title", "]"):
                                continue
                            escaped = re.escape(placeholder)
                            content = re.sub(
                                rf"(<w:t[^>]*>){escaped}(</w:t>)",
                                rf"\1{title_xml}\2",
                                content,
                            )
                        escaped_default = re.escape(_DEFAULT_CLASSIFICATION)
                        content = re.sub(
                            rf"(<w:t[^>]*>){escaped_default}(</w:t>)",
                            rf"\1{classification_xml}\2",
                            content,
                        )
                        with open(filepath, "w", encoding="utf-8") as f:
                            f.write(content)

                with zipfile.ZipFile(docx_path, "w") as out_z:
                    for name in z.namelist():
                        filepath = os.path.join(tmpdir, name)
                        if os.path.exists(filepath):
                            out_z.write(filepath, name)
    finally:
        if os.path.exists(temp_docx):
            os.remove(temp_docx)


def _update_table_fonts(docx_path: Path) -> int:
    """Compact table cell runs to 9pt (python-docx pass)."""
    from docx import Document
    from docx.shared import Pt

    doc = Document(str(docx_path))
    count = 0
    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        count += 1
    doc.save(str(docx_path))
    return count


def apply(
    docx: str | os.PathLike,
    *,
    title: str,
    classification: str = _DEFAULT_CLASSIFICATION,
    author: str = "",
    version: str = "",
    effective_date: str = "",
    code_style: bool = True,
    page_breaks: str = "none",
    font: str = "default",
) -> None:
    """Apply the full SSC post-processing pass to a pandoc-generated DOCX."""
    path = Path(docx)
    _rewrite_zip(path, {"code_style": code_style, "page_breaks": page_breaks,
                        "font": font})
    _update_header(path, title, classification)
    _update_table_fonts(path)
